import os
import time
import threading
from datetime import datetime
import whisper
from docx import Document
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
from openai import OpenAI

# ================= CONFIGURACIÓN =================
AUDIO_FOLDER = r"C:\WhatsApp_Audios"  # Carpeta a vigilar
WORD_FOLDER = os.path.join(AUDIO_FOLDER, "pedidos_word")
os.makedirs(WORD_FOLDER, exist_ok=True)  # crea la carpeta si no existe
MODELO_WHISPER = "small"
API_KEY = "TU_API_KEY"       # Para limpieza de pedido con GPT
MODELO_GPT = "gpt-4.1-mini"
EXTENSION_AUDIO = (".ogg", ".opus", ".mp3", ".wav", ".m4a")
ESPERA_ARCHIVO = 2        # Segundos entre intentos de espera
REINTENTOS_WHISPER = 3    # Reintentos de Whisper
TEXTO_MINIMO = 10         # Longitud mínima de transcripción
MAX_ESPERA = 10           # Max intentos para esperar archivo listo
# ================================================

print("🔄 Cargando Whisper...")
whisper_model = whisper.load_model(MODELO_WHISPER)
client = OpenAI(api_key=API_KEY)
print("✅ Sistema listo. Solo procesará audios nuevos.\n")

# ---------------- FUNCIONES ----------------
def generar_nombre_word(nombre_original):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    safe_nombre = os.path.splitext(os.path.basename(nombre_original))[0]
    safe_nombre = safe_nombre.replace(" ", "_").replace("(", "_").replace(")", "_")
    return os.path.join(WORD_FOLDER, f"Pedido_{safe_nombre}_{timestamp}.docx")

def guardar_word(texto, nombre_original):
    doc = Document()
    section = doc.sections[0]
    header = section.header
    if header.is_linked_to_previous:
        header.is_linked_to_previous = False
    if not header.paragraphs:
        header_para = header.add_paragraph()
    else:
        header_para = header.paragraphs[0]
    header_para.text = "Pedido de Cliente"
    for linea in texto.splitlines():
        doc.add_paragraph(linea)
    nombre_archivo = generar_nombre_word(nombre_original)
    doc.save(nombre_archivo)
    print(f"✅ Pedido generado: {os.path.basename(nombre_archivo)}\n")

def esperar_archivo_listo(ruta_audio):
    """Espera hasta que el archivo tenga tamaño >0 y sea estable"""
    intentos = 0
    tamaño_prev = -1
    while intentos < MAX_ESPERA:
        if os.path.exists(ruta_audio):
            tamaño_actual = os.path.getsize(ruta_audio)
            if tamaño_actual > 0 and tamaño_actual == tamaño_prev:
                return True
            tamaño_prev = tamaño_actual
        time.sleep(ESPERA_ARCHIVO)
        intentos += 1
    return False

def transcribir_audio(ruta_audio):
    for intento in range(1, REINTENTOS_WHISPER+1):
        try:
            print(f"🎧 Transcribiendo: {os.path.basename(ruta_audio)} (Intento {intento})")
            result = whisper_model.transcribe(ruta_audio, language="es")
            texto = result["text"].strip()
            if len(texto) < TEXTO_MINIMO:
                print(f"⚠️ Texto muy corto, se ignora: {os.path.basename(ruta_audio)}")
                return None
            return texto
        except Exception as e:
            print(f"⚠️ Error transcribiendo {os.path.basename(ruta_audio)}: {e}")
            time.sleep(ESPERA_ARCHIVO)
    print(f"❌ Fallaron todos los intentos para {os.path.basename(ruta_audio)}")
    return None

def limpiar_pedido_con_gpt(texto):
    print("🧠 Limpiando pedido con IA...")
    response = client.chat.completions.create(
        model=MODELO_GPT,
        temperature=0,
        messages=[
            {"role": "system", "content": "Eres un asistente que limpia pedidos de restaurante."},
            {"role": "user", "content": f"""
Limpia y ordena este pedido.

Reglas:
- Devuelve SOLO un listado simple
- Una línea por producto
- Cantidad + producto
- Sin saludos ni frases
- Español claro para proveedor

Texto:
{texto}
"""}
        ]
    )
    return response.choices[0].message.content.strip()

def procesar_pedido(ruta_audio):
    try:
        # Ignorar duplicados temporales
        if "(" in os.path.basename(ruta_audio) and ")" in os.path.basename(ruta_audio):
            return
        if not esperar_archivo_listo(ruta_audio):
            print(f"⚠️ Archivo no listo o vacío: {os.path.basename(ruta_audio)}")
            return

        texto_sucio = transcribir_audio(ruta_audio)
        if not texto_sucio:
            return

        # Guardar texto original
        with open(os.path.join(AUDIO_FOLDER, "pedido_sucio.txt"), "w", encoding="utf-8") as f:
            f.write(texto_sucio)

        # Limpiar con GPT
        texto_limpio = limpiar_pedido_con_gpt(texto_sucio)

        # Guardar texto limpio
        with open(os.path.join(AUDIO_FOLDER, "pedido_limpio.txt"), "w", encoding="utf-8") as f:
            f.write(texto_limpio)

        # Crear Word
        guardar_word(texto_limpio, ruta_audio)

    except Exception as e:
        print(f"❌ Error procesando {os.path.basename(ruta_audio)}: {e}\n")

# ---------------- WATCHDOG ----------------
class AudioHandler(FileSystemEventHandler):
    def on_created(self, event):
        if event.is_directory:
            return
        ruta_audio = event.src_path
        if not ruta_audio.lower().endswith(EXTENSION_AUDIO):
            return

        # Hilo independiente con espera hasta que archivo esté listo
        threading.Thread(target=procesar_pedido, args=(ruta_audio,), daemon=True).start()

# ---------------- EJECUCIÓN ----------------
observer = Observer()
observer.schedule(AudioHandler(), AUDIO_FOLDER, recursive=False)
observer.start()
print(f"[✔] Vigilando la carpeta: {AUDIO_FOLDER}...\n")

try:
    while True:
        time.sleep(1)
except KeyboardInterrupt:
    observer.stop()
observer.join()